import asyncio
import os
import shutil
import sys
import tempfile
import time
from pathlib import Path

# Add backend root to path
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.base import AsyncSessionLocal
from app.models.fir import FIR, FileType
from app.models.fir_document_content import FIRDocumentContent
from app.db.init_db import check_db_connection

# Start uvicorn server in background
import subprocess

async def run_tests():
    print("=== STARTING EXTRACTION TEST SUITE ===")
    
    # 1. Ensure DB connection & tables are initialized
    print("\n[Test DB] Initializing database and tables...")
    db_ok = await check_db_connection()
    if not db_ok:
        print("❌ Database connection check failed.")
        sys.exit(1)
    print("✅ Database connection check passed.")

    # Check if a server is already running on port 8000
    server_already_running = False
    client = httpx.AsyncClient(base_url="http://127.0.0.1:8000", timeout=30.0)
    try:
        res = await client.get("/api/v1/health")
        if res.status_code == 200:
            print("🚀 Detected that a FastAPI backend server is already running on port 8000. Reusing it.")
            server_already_running = True
    except Exception:
        pass

    server_process = None
    if not server_already_running:
        # Start uvicorn backend server on port 8000
        print("\nStarting FastAPI server...")
        test_env = os.environ.copy()
        test_env["ENKRYPT_ENABLED"] = "false"
        test_env["PYTHONUNBUFFERED"] = "1"
        server_process = subprocess.Popen(
            [".venv/bin/python", "-u", "-m", "uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", "8000"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=test_env
        )
        
        # Wait for server to boot
        time.sleep(5)
        
        # Check if the process exited immediately
        poll_status = server_process.poll()
        if poll_status is not None:
            print(f"❌ Server process exited immediately with code {poll_status}")
            stdout, stderr = server_process.communicate()
            print("--- SERVER STDOUT ---")
            print(stdout)
            print("--- SERVER STDERR ---")
            print(stderr)
            sys.exit(1)
            
        # Verify health
        try:
            res = await client.get("/api/v1/health")
            print(f"Server health check: {res.status_code} {res.json()}")
        except Exception as e:
            print(f"❌ Server failed to start: {e}")
            # Terminate first so communicate() returns immediately
            server_process.terminate()
            stdout, stderr = server_process.communicate()
            print("--- SERVER STDOUT ---")
            print(stdout)
            print("--- SERVER STDERR ---")
            print(stderr)
            sys.exit(1)

    # Prepare real and corrupted test files
    temp_dir = tempfile.mkdtemp()
    
    # 1. TXT File
    txt_path = Path(temp_dir) / "test_fir.txt"
    txt_path.write_text("This is a plain text First Information Report (FIR) from Karnataka Police. The suspect was seen near Vidhana Soudha at 10:00 PM.")
    
    # 2. PDF File (Generated programmatically using PyMuPDF if installed, fallback to plain text fake)
    pdf_path = Path(temp_dir) / "test_fir.pdf"
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "This is a valid PDF document with text content for test cases.")
        doc.save(str(pdf_path))
        doc.close()
        print("Generated valid PDF using PyMuPDF.")
    except Exception as e:
        print(f"Could not use PyMuPDF to generate valid test PDF ({e}). Writing mock pdf.")
        pdf_path.write_text("mock pdf data")

    # 3. DOCX File (Generated using python-docx)
    docx_path = Path(temp_dir) / "test_fir.docx"
    try:
        import docx
        doc = docx.Document()
        doc.add_paragraph("This is a valid Word Document (DOCX) for FIR extraction testing.")
        doc.add_paragraph("Witness reported seeing the suspect at Majestic bus station.")
        doc.save(str(docx_path))
        print("Generated valid DOCX using python-docx.")
    except Exception as e:
        print(f"Could not use python-docx to generate valid test DOCX ({e}). Writing mock docx.")
        docx_path.write_text("mock docx data")

    # 4. Corrupted PDF File
    corrupted_pdf_path = Path(temp_dir) / "corrupted.pdf"
    corrupted_pdf_path.write_bytes(b"INVALID PDF DATA CHUNKS")

    # 5. Corrupted DOCX File
    corrupted_docx_path = Path(temp_dir) / "corrupted.docx"
    corrupted_docx_path.write_bytes(b"INVALID DOCX DATA CHUNKS")

    # 6. Empty Document
    empty_txt_path = Path(temp_dir) / "empty.txt"
    empty_txt_path.write_text("")

    uploaded_firs = {}

    try:
        # Helper function to upload and extract
        async def upload_file(path: Path, case_number: str, mime_type: str):
            with open(path, "rb") as f:
                res = await client.post(
                    "/api/v1/firs/upload",
                    data={"case_number": case_number, "created_by": "tester"},
                    files={"file": (path.name, f, mime_type)}
                )
            if res.status_code != 201:
                raise ValueError(f"Upload failed: {res.status_code} {res.text}")
            return res.json()["id"]

        # Upload files
        print("\nUploading test documents...")
        uploaded_firs["txt"] = await upload_file(txt_path, "CASE-TXT-101", "text/plain")
        uploaded_firs["pdf"] = await upload_file(pdf_path, "CASE-PDF-102", "application/pdf")
        uploaded_firs["docx"] = await upload_file(docx_path, "CASE-DOCX-103", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        uploaded_firs["empty"] = await upload_file(empty_txt_path, "CASE-EMPTY-104", "text/plain")
        uploaded_firs["corr_pdf"] = await upload_file(corrupted_pdf_path, "CASE-CORR-PDF-105", "application/pdf")
        uploaded_firs["corr_docx"] = await upload_file(corrupted_docx_path, "CASE-CORR-DOCX-106", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        print("✅ All test documents uploaded successfully.")

        # Test 1: PDF Extraction
        print("\n[Test 1] Extracting PDF Document...")
        pdf_id = uploaded_firs["pdf"]
        res = await client.post(f"/api/v1/firs/{pdf_id}/extract")
        assert res.status_code == 200, f"PDF Extraction failed: {res.status_code} {res.text}"
        data = res.json()
        print(f"✅ PDF Extracted: {data}")
        assert "valid PDF document" in data["extracted_text"]
        assert data["page_count"] == 1
        assert data["word_count"] > 0
        assert data["character_count"] > 0
        assert data["language"] == "en"

        # Test 2: DOCX Extraction
        print("\n[Test 2] Extracting DOCX Document...")
        docx_id = uploaded_firs["docx"]
        res = await client.post(f"/api/v1/firs/{docx_id}/extract")
        assert res.status_code == 200, f"DOCX Extraction failed: {res.status_code} {res.text}"
        data = res.json()
        print(f"✅ DOCX Extracted: {data}")
        assert "valid Word Document" in data["extracted_text"]
        assert data["page_count"] >= 1
        assert data["word_count"] > 0
        assert data["character_count"] > 0

        # Test 3: TXT Extraction
        print("\n[Test 3] Extracting TXT Document...")
        txt_id = uploaded_firs["txt"]
        res = await client.post(f"/api/v1/firs/{txt_id}/extract")
        assert res.status_code == 200, f"TXT Extraction failed: {res.status_code} {res.text}"
        data = res.json()
        print(f"✅ TXT Extracted: {data}")
        assert "plain text First Information Report" in data["extracted_text"]
        assert data["word_count"] > 0
        assert data["character_count"] > 0

        # Test 4: Empty Document
        print("\n[Test 4] Extracting Empty Document...")
        empty_id = uploaded_firs["empty"]
        res = await client.post(f"/api/v1/firs/{empty_id}/extract")
        assert res.status_code == 200, f"Empty document extraction failed: {res.status_code} {res.text}"
        data = res.json()
        print(f"✅ Empty Document Extracted: {data}")
        assert data["extracted_text"] == ""
        assert data["word_count"] == 0
        assert data["character_count"] == 0

        # Test 5: Corrupted PDF Document (should fail)
        print("\n[Test 5] Extracting Corrupted PDF...")
        corr_pdf_id = uploaded_firs["corr_pdf"]
        res = await client.post(f"/api/v1/firs/{corr_pdf_id}/extract")
        assert res.status_code == 422, f"Expected 422 for corrupted PDF, got: {res.status_code} {res.text}"
        print(f"✅ Corrupted PDF correctly blocked: {res.json()['detail']}")

        # Test 6: Corrupted DOCX Document (should fail)
        print("\n[Test 6] Extracting Corrupted DOCX...")
        corr_docx_id = uploaded_firs["corr_docx"]
        res = await client.post(f"/api/v1/firs/{corr_docx_id}/extract")
        assert res.status_code == 422, f"Expected 422 for corrupted DOCX, got: {res.status_code} {res.text}"
        print(f"✅ Corrupted DOCX correctly blocked: {res.json()['detail']}")

        # Test 7: Duplicate extraction updates existing record instead of creating new ones
        print("\n[Test 7] Duplicate Extraction (Update)...")
        res1 = await client.post(f"/api/v1/firs/{txt_id}/extract")
        assert res1.status_code == 200
        data1 = res1.json()
        
        # Trigger second extraction
        res2 = await client.post(f"/api/v1/firs/{txt_id}/extract")
        assert res2.status_code == 200
        data2 = res2.json()
        
        assert data1["id"] == data2["id"], "ID changed during duplicate extraction update!"
        print(f"✅ Duplicate extraction returns same record ID: {data2['id']}")

        # Test 8: Verify database records exist
        print("\n[Test 8] Querying Database directly to verify records...")
        async with AsyncSessionLocal() as session:
            # Check total records in fir_document_contents
            stmt = select(FIRDocumentContent)
            res_db = await session.execute(stmt)
            records = res_db.scalars().all()
            print(f"✅ Found {len(records)} extracted document content records in PostgreSQL.")
            assert len(records) == 4, f"Expected exactly 4 records, got {len(records)}"
            
            # Verify one specific record has preview / text
            stmt_specific = select(FIRDocumentContent).where(FIRDocumentContent.fir_id == txt_id)
            res_specific = await session.execute(stmt_specific)
            spec_rec = res_specific.scalar_one_or_none()
            assert spec_rec is not None
            print(f"✅ Database record verified: Preview (first 1000 chars):\n'{spec_rec.extracted_text[:100]}...'")

    finally:
        # Cleanup uploaded files from DB
        print("\nCleaning up uploaded files...")
        async with AsyncSessionLocal() as session:
            for key, fid in uploaded_firs.items():
                try:
                    # Retrieve storage path
                    res_fir = await session.execute(select(FIR).where(FIR.id == fid))
                    fir_rec = res_fir.scalar_one_or_none()
                    if fir_rec:
                        storage_path = Path(fir_rec.storage_path)
                        if storage_path.exists():
                            storage_path.unlink()
                        await session.delete(fir_rec)
                except Exception as e:
                    print(f"Error during db cleanup for {fid}: {e}")
            await session.commit()
            
        shutil.rmtree(temp_dir)
        if server_process is not None:
            server_process.terminate()
            server_process.wait()
        await client.aclose()

    print("\n=== ALL EXTRACTION TESTS PASSED SUCCESSFULLY! ===")

if __name__ == "__main__":
    asyncio.run(run_tests())
