"""
app/services/document_processor.py
──────────────────────────────────
Service for extracting text and metadata from PDF, DOCX, and TXT files.
"""

import os
import logging
from typing import TypedDict

# Configure logging
logger = logging.getLogger(__name__)

# Try to import langdetect for language detection
try:
    import langdetect
    HAS_LANGDETECT = True
except ImportError:
    HAS_LANGDETECT = False
    logger.warning("langdetect library not found, language detection will fallback to 'en'.")


class ExtractionResult(TypedDict):
    text: str
    page_count: int
    word_count: int
    character_count: int
    language: str


def detect_language(text: str) -> str:
    """Helper to detect language of the extracted text with fallback to 'en'."""
    if not text.strip():
        return "en"
    if not HAS_LANGDETECT:
        return "en"
    try:
        # Detect language using first 1000 characters for performance and accuracy
        lang = langdetect.detect(text[:1000])
        return lang if lang else "en"
    except Exception as e:
        logger.debug(f"Language detection failed, fallback to 'en': {e}")
        return "en"


def extract_text_from_pdf(file_path: str) -> ExtractionResult:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF (fitz) is not installed on the system.")

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error(f"Failed to open PDF file {file_path}: {e}")
        raise ValueError(f"Corrupted or invalid PDF file: {e}")

    text_parts = []
    page_count = len(doc)
    
    for page_idx in range(page_count):
        try:
            page = doc[page_idx]
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"Error extracting text from page {page_idx} of PDF {file_path}: {e}")

    text = "\n".join(text_parts)
    char_count = len(text)
    word_count = len(text.split())
    language = detect_language(text)

    # Ensure page count is at least 1 if document contains pages, or 0 if empty
    resolved_page_count = page_count if page_count > 0 else 0

    return {
        "text": text,
        "page_count": resolved_page_count,
        "word_count": word_count,
        "character_count": char_count,
        "language": language,
    }


def extract_text_from_docx(file_path: str) -> ExtractionResult:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is not installed on the system.")

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logger.error(f"Failed to open DOCX file {file_path}: {e}")
        raise ValueError(f"Corrupted or invalid DOCX file: {e}")

    text_parts = []
    for para in doc.paragraphs:
        if para.text:
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text_parts.append(" | ".join(row_text))

    text = "\n".join(text_parts)
    char_count = len(text)
    word_count = len(text.split())
    language = detect_language(text)

    # DOCX doesn't have a direct physical page count property.
    # We estimate based on average character count (approx 2000 chars per page).
    estimated_pages = max(1, char_count // 2000) if char_count > 0 else 0

    return {
        "text": text,
        "page_count": estimated_pages,
        "word_count": word_count,
        "character_count": char_count,
        "language": language,
    }


def extract_text_from_txt(file_path: str) -> ExtractionResult:
    """Extract text from a plain TXT file."""
    # Attempt to read the file using UTF-8, fallback to Latin-1 if it fails
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        except Exception as e:
            logger.error(f"Failed to read TXT file {file_path} with Latin-1: {e}")
            raise ValueError(f"Unable to read text file: {e}")
    except Exception as e:
        logger.error(f"Failed to open TXT file {file_path}: {e}")
        raise ValueError(f"Unable to open text file: {e}")

    char_count = len(text)
    word_count = len(text.split())
    language = detect_language(text)

    # TXT files don't have pages. Estimate 2000 characters per page.
    estimated_pages = max(1, char_count // 2000) if char_count > 0 else 0

    return {
        "text": text,
        "page_count": estimated_pages,
        "word_count": word_count,
        "character_count": char_count,
        "language": language,
    }


def extract_text(file_path: str) -> ExtractionResult:
    """
    Automatically detects the file extension and extracts text.
    Raises ValueError for corrupted files or unsupported extensions.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found on disk: {file_path}")

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in [".txt", ".text"]:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format for text extraction: '{ext}'")
